from __future__ import annotations

import re
import shutil
import uuid
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

import yaml

from ..config import Settings


MAX_UPLOAD_BYTES = 25 * 1024 * 1024
MAX_EXTRACTED_CHARS = 500_000
MAX_PDF_PAGES = 80
MAX_DOCX_PARAGRAPHS = 20_000
MAX_DOCX_TABLE_CELLS = 100_000
SUPPORTED_UPLOAD_SUFFIXES = {".md", ".txt", ".pdf", ".docx"}
WIKI_ROOTS = {"public", "private"}


@dataclass(frozen=True)
class UploadedFile:
    filename: str
    content: bytes


@dataclass(frozen=True)
class PreparedWikiUpload:
    original_name: str
    suffix: str
    content: str


@dataclass(frozen=True)
class StagedWikiUpload:
    item: PreparedWikiUpload
    staging_path: Path
    output_name: str


def wiki_roots(workspace: Path) -> dict[str, Path]:
    settings = Settings(_env_file=str(workspace / ".env") if (workspace / ".env").exists() else None)
    return {
        "public": _resolve_under_workspace(workspace, settings.WIKI_DIR),
        "private": _resolve_under_workspace(workspace, settings.WIKI_AUTO_RECORD_DIR),
    }


def build_wiki_tree(workspace: Path) -> dict[str, Any]:
    roots = wiki_roots(workspace)
    payload: dict[str, Any] = {"roots": {}}
    for key, root in roots.items():
        tree = _directory_node(root, root) if root.is_dir() else _empty_directory(root.name or key)
        payload["roots"][key] = {
            "key": key,
            "label": "公开知识库" if key == "public" else "私有知识库",
            "path": str(root),
            "exists": root.is_dir(),
            "file_count": tree["total_files"],
            "dir_count": _count_dirs(tree),
            "tree": tree,
        }
    return payload


def read_wiki_file(workspace: Path, root_key: str, rel_path: str) -> dict[str, Any]:
    root = _root_for_key(workspace, root_key)
    path = _resolve_wiki_file(root, rel_path)
    stat = path.stat()
    return {
        "root": root_key,
        "path": path.relative_to(root).as_posix(),
        "name": path.name,
        "size": stat.st_size,
        "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "content": path.read_text(encoding="utf-8", errors="replace"),
    }


def import_uploaded_files(
    workspace: Path,
    files: list[UploadedFile],
    *,
    root_key: str = "private",
) -> dict[str, Any]:
    if not files:
        raise ValueError("No files uploaded")
    root = _root_for_key(workspace, root_key)
    upload_dir = root / "uploads"

    prepared = [_prepare_upload(item) for item in files]
    upload_dir.mkdir(parents=True, exist_ok=True)
    staging_dir = upload_dir / f".staging-{uuid.uuid4().hex}"
    committed_paths: list[Path] = []
    try:
        staged = _stage_prepared_uploads(upload_dir, staging_dir, prepared)
        imported = _commit_staged_uploads(upload_dir, staged, committed_paths)
        return {"uploaded": imported, "tree": build_wiki_tree(workspace)}
    except Exception:
        _rollback_committed_uploads(committed_paths)
        raise
    finally:
        shutil.rmtree(staging_dir, ignore_errors=True)


def _prepare_upload(item: UploadedFile) -> PreparedWikiUpload:
    original_name = Path(item.filename or "").name
    if not original_name:
        raise ValueError("Uploaded file must have a filename")
    if len(item.content) == 0:
        raise ValueError(f"{original_name} is empty")
    if len(item.content) > MAX_UPLOAD_BYTES:
        raise ValueError(f"{original_name} exceeds {MAX_UPLOAD_BYTES // (1024 * 1024)} MB")

    suffix = Path(original_name).suffix.lower()
    if suffix not in SUPPORTED_UPLOAD_SUFFIXES:
        allowed = ", ".join(sorted(SUPPORTED_UPLOAD_SUFFIXES))
        raise ValueError(f"{original_name} is not supported. Allowed: {allowed}")

    body = _extract_markdown_body(original_name, item.content, suffix)
    content = _with_upload_frontmatter(original_name, suffix, body)
    return PreparedWikiUpload(original_name=original_name, suffix=suffix, content=content)


def _stage_prepared_uploads(
    upload_dir: Path,
    staging_dir: Path,
    prepared: list[PreparedWikiUpload],
) -> list[StagedWikiUpload]:
    staging_dir.mkdir(parents=True, exist_ok=False)
    reserved_names: set[str] = set()
    staged: list[StagedWikiUpload] = []
    for item in prepared:
        output_name = _unique_markdown_name(upload_dir, item.original_name, reserved_names)
        reserved_names.add(output_name)
        staging_path = staging_dir / output_name
        staging_path.write_text(item.content, encoding="utf-8")
        staged.append(
            StagedWikiUpload(item=item, staging_path=staging_path, output_name=output_name)
        )
    return staged


def _commit_staged_uploads(
    upload_dir: Path,
    staged: list[StagedWikiUpload],
    committed_paths: list[Path],
) -> list[dict[str, Any]]:
    imported: list[dict[str, Any]] = []
    for item in staged:
        output_path = upload_dir / item.output_name
        _commit_staged_upload(item.staging_path, output_path)
        committed_paths.append(output_path)
        imported.append(_upload_result(upload_dir, item.item, output_path))
    return imported


def _commit_staged_upload(staging_path: Path, output_path: Path) -> None:
    if output_path.exists():
        raise FileExistsError(f"Upload target already exists: {output_path.name}")
    staging_path.replace(output_path)


def _rollback_committed_uploads(paths: list[Path]) -> None:
    for path in reversed(paths):
        try:
            if path.is_file():
                path.unlink()
        except OSError:
            pass


def _upload_result(
    upload_dir: Path,
    item: PreparedWikiUpload,
    output_path: Path,
) -> dict[str, Any]:
    return {
        "original_name": item.original_name,
        "path": output_path.relative_to(upload_dir.parent).as_posix(),
        "size": output_path.stat().st_size,
        "source_type": item.suffix.lstrip("."),
    }


def _extract_markdown_body(original_name: str, content: bytes, suffix: str) -> str:
    if suffix == ".md":
        return _decode_text(content)
    if suffix == ".txt":
        return f"# {Path(original_name).stem}\n\n{_decode_text(content)}"
    if suffix == ".pdf":
        return _extract_pdf_text(original_name, content)
    if suffix == ".docx":
        return _extract_docx_text(original_name, content)
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf_text(original_name: str, content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ValueError("PDF upload requires pypdf") from exc
    try:
        reader = PdfReader(BytesIO(content))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError(f"{original_name} has too many pages; limit is {MAX_PDF_PAGES}")
        parts = []
        total_chars = 0
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                total_chars = _append_limited(
                    parts,
                    f"## Page {index}\n\n{text}",
                    total_chars,
                    original_name,
                )
    except Exception as exc:
        raise ValueError(f"Cannot extract text from {original_name}: {exc}") from exc
    if not parts:
        raise ValueError(f"Cannot extract text from {original_name}: no text found")
    return f"# {Path(original_name).stem}\n\n" + "\n\n".join(parts)


def _extract_docx_text(original_name: str, content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ValueError("DOCX upload requires python-docx") from exc
    try:
        doc = Document(BytesIO(content))
        if len(doc.paragraphs) > MAX_DOCX_PARAGRAPHS:
            raise ValueError(
                f"{original_name} has too many paragraphs; limit is {MAX_DOCX_PARAGRAPHS}"
            )
        parts = []
        total_chars = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                total_chars = _append_limited(parts, text, total_chars, original_name)
        table_cells = 0
        for table in doc.tables:
            for row in table.rows:
                table_cells += len(row.cells)
                if table_cells > MAX_DOCX_TABLE_CELLS:
                    raise ValueError(
                        f"{original_name} has too many table cells; "
                        f"limit is {MAX_DOCX_TABLE_CELLS}"
                    )
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    total_chars = _append_limited(
                        parts,
                        " | ".join(cells),
                        total_chars,
                        original_name,
                    )
    except Exception as exc:
        raise ValueError(f"Cannot extract text from {original_name}: {exc}") from exc
    if not parts:
        raise ValueError(f"Cannot extract text from {original_name}: no text found")
    return f"# {Path(original_name).stem}\n\n" + "\n\n".join(parts)


def _append_limited(
    parts: list[str],
    chunk: str,
    total_chars: int,
    original_name: str,
) -> int:
    next_total = total_chars + len(chunk)
    if next_total > MAX_EXTRACTED_CHARS:
        raise ValueError(
            f"{original_name} extracted text exceeds {MAX_EXTRACTED_CHARS} characters"
        )
    parts.append(chunk)
    return next_total


def _with_upload_frontmatter(original_name: str, suffix: str, body: str) -> str:
    frontmatter = {
        "title": Path(original_name).stem,
        "type": "concept",
        "tags": ["uploaded", suffix.lstrip(".")],
        "sources": [],
        "uploaded_at": datetime.now().isoformat(timespec="seconds"),
        "original_filename": original_name,
        "source_type": suffix.lstrip("."),
    }
    front = yaml.safe_dump(frontmatter, sort_keys=False, allow_unicode=True)
    return f"---\n{front}---\n\n{body.strip()}\n"


def _directory_node(root: Path, path: Path) -> dict[str, Any]:
    rel = "" if path == root else path.relative_to(root).as_posix()
    children = []
    direct_files = 0
    total_files = 0
    for child in sorted(path.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower())):
        if child.name.startswith(".") or child.is_symlink():
            continue
        if child.is_dir():
            node = _directory_node(root, child)
            if node["total_files"] > 0:
                total_files += node["total_files"]
                children.append(node)
        elif child.is_file() and child.suffix.lower() == ".md":
            direct_files += 1
            total_files += 1
            stat = child.stat()
            children.append(
                {
                    "type": "file",
                    "name": child.name,
                    "path": child.relative_to(root).as_posix(),
                    "size": stat.st_size,
                    "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(
                        timespec="seconds"
                    ),
                }
            )
    return {
        "type": "directory",
        "name": path.name or str(path),
        "path": rel,
        "file_count": direct_files,
        "total_files": total_files,
        "children": children,
    }


def _empty_directory(name: str) -> dict[str, Any]:
    return {
        "type": "directory",
        "name": name,
        "path": "",
        "file_count": 0,
        "total_files": 0,
        "children": [],
    }


def _count_dirs(node: dict[str, Any]) -> int:
    if node.get("type") != "directory":
        return 0
    return 1 + sum(_count_dirs(child) for child in node.get("children", []))


def _root_for_key(workspace: Path, root_key: str) -> Path:
    if root_key not in WIKI_ROOTS:
        raise ValueError("root must be public or private")
    return wiki_roots(workspace)[root_key]


def _resolve_wiki_file(root: Path, rel_path: str) -> Path:
    if not rel_path or Path(rel_path).is_absolute():
        raise ValueError("Invalid wiki path")
    raw = Path(rel_path.replace("\\", "/"))
    if any(part in {"", ".", ".."} or part.startswith(".") for part in raw.parts):
        raise ValueError("Invalid wiki path")
    path = (root / raw).resolve()
    root_resolved = root.resolve()
    try:
        path.relative_to(root_resolved)
    except ValueError as exc:
        raise ValueError("Invalid wiki path") from exc
    if path.suffix.lower() != ".md":
        raise ValueError("Only Markdown files can be read")
    if not path.is_file():
        raise ValueError("Wiki file not found")
    return path


def _resolve_under_workspace(workspace: Path, configured: str) -> Path:
    workspace_resolved = workspace.resolve()
    path = Path(configured)
    if not path.is_absolute():
        path = workspace / path
    resolved = path.resolve()
    try:
        resolved.relative_to(workspace_resolved)
    except ValueError as exc:
        raise ValueError("Wiki directories must stay under the workspace") from exc
    return resolved


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _unique_markdown_name(
    upload_dir: Path,
    original_name: str,
    reserved_names: set[str] | None = None,
) -> str:
    reserved_names = reserved_names or set()
    stem = _safe_stem(Path(original_name).stem)
    today = datetime.now().strftime("%Y%m%d-%H%M%S")
    candidate = f"{today}-{stem}.md"
    if candidate not in reserved_names and not (upload_dir / candidate).exists():
        return candidate
    index = 2
    while True:
        candidate = f"{today}-{stem}-{index}.md"
        if candidate not in reserved_names and not (upload_dir / candidate).exists():
            return candidate
        index += 1


def _safe_stem(value: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip(".-_")
    stem = re.sub(r"-{2,}", "-", stem)
    return (stem or "upload")[:80]
